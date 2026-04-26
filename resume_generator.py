import os
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
import json
import re

load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# ✅ UPDATED FUNCTION (NOW USES ANALYSIS)
def generate_resume(job_description, resume_text, analysis_text):

    # Clean analysis JSON (remove ``` if present)
    cleaned = analysis_text.strip()
    cleaned = re.sub(r"```json|```", "", cleaned).strip()

    prompt = f"""
You are an expert resume writer for senior enterprise roles.

You are given:

1. Candidate Resume
2. Job Description
3. Gap Analysis (from recruiter perspective)

------------------------
Candidate Resume:
{resume_text}

------------------------
Job Description:
{job_description}

------------------------
Gap Analysis:
{cleaned}

------------------------

Your task:

- Improve the resume using the gap analysis
- Add missing keywords ONLY if they are realistically inferable
- Do NOT invent fake experience
- Rephrase existing experience to better match job requirements
- Strengthen:
  - Salesforce Release Management
  - DevOps / CI-CD
  - SFDX / Metadata Deployment
  - Governance / Documentation
- Keep executive tone
- ATS optimized
- Keep structure:

  Name & Title  
  Summary  
  Core Skills  
  Professional Experience  
  Certifications  

IMPORTANT:
- Do NOT add fake tools or certifications
- If something is implied, make it explicit
- Make resume stronger but truthful

Return clean formatted resume text.
"""

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[{"role": "user", "content": prompt}]
    )

    return response.choices[0].message.content


# ✅ WORD EXPORT (UNCHANGED)
def save_resume_to_word(content, filename="Tailored_Resume.docx"):
    doc = Document()

    for line in content.split("\n"):
        line = line.strip()

        if not line:
            continue

        if line.isupper() or line.endswith(":"):
            doc.add_heading(line, level=1)

        elif line.startswith("-"):
            doc.add_paragraph(line[1:].strip(), style="List Bullet")

        else:
            doc.add_paragraph(line)

    doc.save(filename)
    return filename